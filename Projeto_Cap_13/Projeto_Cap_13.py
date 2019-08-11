"""
Created on Sat May 19 17:25:06 2018

@author: Edson Frota
"""

import docx
doc = docx.Document()
doc.add_heading('Convite de Aniversário',0)
doc.add_heading('Venha Participar da Minha Festa, Você é meu convidado, traga sua família!')
doc.add_heading('Local: Casa da Vóvó')
doc.add_heading('Horário: 20:00h')
doc.add_heading('Traga Presente')
linha_01 = doc.add_paragraph('Sugestão de presente: ')
linha_01.add_run('Brinquedos!, Eletronicos!, Comida!')

doc.add_picture('baloes.png', width=docx.shared.Inches(8),height=docx.shared.Cm(8))

doc.save('Doc_cap_13.docx')
